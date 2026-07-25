import os
import json
import re
import time
import datetime
import traceback
import psycopg2
import pandas as pd
import plotly.express as px
from openai import OpenAI
import requests
from PIL import Image
import markdown
import ssl

from db import get_db
from models import (
    SystemConfig,
    ConnectionSetting,
    Prompt,
    SavedQuery,
    ReportPromptCase,
    LLMProvider,
    LLMModel,
    LLMFallback,
)

# =========================
# СИСТЕМНЫЙ КОНФИГ (ранее config.txt)
# =========================
def _load_system_config():
    keys = [
        "redash",
        "redash_url",
        "redash_datasource_id",
        "bothub",
        "LLM",
        "LLM2",
        "LLM3",
        "LLM4",
        "LLM5",
        "check",
        "download_dir",
        "query_limit",
        "db_desc",
    ]
    data = {}
    with get_db() as db:
        for key in keys:
            row = db.query(SystemConfig).filter(SystemConfig.key == key).first()
            data[key] = row.value if row else ""
    return data


def get_system_config():
    try:
        return _load_system_config()
    except Exception:
        return {}


# Все константы теперь вычисляются на лету
def get_redash_config():
    try:
        cfg = get_system_config()
        api_key = cfg.get("redash", "").strip()
        url = cfg.get("redash_url", "http://dash_server:5000/api/queries").strip()
        try:
            datasource_id = int(cfg.get("redash_datasource_id", 1))
        except (TypeError, ValueError):
            datasource_id = 1
        return api_key, url, datasource_id
    except Exception:
        return "", "http://dash_server:5000/api/queries", 1

REDASH_API_KEY, REDASH_URL, REDASH_DATA_SOURCE_ID = get_redash_config()

def get_query_limit():
    cfg = get_system_config()
    try:
        return max(100, int(cfg.get("query_limit", 1000)))
    except (TypeError, ValueError):
        return 1000

QUERY_LIMIT = get_query_limit()

def get_llm_token():
    cfg = get_system_config()
    return cfg.get("bothub", "").strip()

LLM_TOKEN = get_llm_token()

def get_llm_model():
    cfg = get_system_config()
    return cfg.get("LLM", "gpt-4.1-nano").strip()

LLM_MODEL = get_llm_model()

def get_db_description():
    cfg = get_system_config()
    desc = cfg.get("db_desc", "")
    if desc:
        return desc
    return ""

def get_check_sql():
    cfg = get_system_config()
    check_raw = cfg.get("check", "yes").strip().lower()
    return check_raw in ("yes", "true", "1")

check_sql = get_check_sql()

# Удалены все кеширующие переменные: _system_config_cache, _llm_config_cache, _schema_cache

def refresh_system_config():
    pass

def get_default_llm_config():
    with get_db() as db:
        model = db.query(LLMModel).filter(LLMModel.is_default == True, LLMModel.enabled == True).first()
        if model:
            provider = db.query(LLMProvider).filter(LLMProvider.id == model.provider_id, LLMProvider.enabled == True).first()
            if provider:
                return {
                    "model_id": str(model.id),
                    "model_name": model.model_name,
                    "api_key": provider.api_key,
                    "base_url": provider.base_url,
                    "temperature": float(model.temperature) if model.temperature else None,
                    "timeout": model.timeout or 120,
                }
    
    log_message("LLM_FALLBACK", "Нет активной модели с is_default=True. Использую fallback: LLM_MODEL")
    return {
        "model_id": None,
        "model_name": LLM_MODEL,
        "api_key": LLM_TOKEN,
        "base_url": "https://bothub.chat/api/v2/openai/v1",
        "temperature": None,
        "timeout": 120,
    }


def get_llm_config_for_model(model_id):
    with get_db() as db:
        model = db.query(LLMModel).filter(LLMModel.id == model_id, LLMModel.enabled == True).first()
        if not model:
            return None
        provider = db.query(LLMProvider).filter(LLMProvider.id == model.provider_id, LLMProvider.enabled == True).first()
        if not provider:
            return None
        return {
            "model_id": str(model.id),
            "model_name": model.model_name,
            "api_key": provider.api_key,
            "base_url": provider.base_url,
            "temperature": float(model.temperature) if model.temperature else None,
            "timeout": model.timeout or 120,
        }


def get_fallback_configs(model_id):
    configs = []
    with get_db() as db:
        fallbacks = db.query(LLMFallback).filter(LLMFallback.model_id == model_id).order_by(LLMFallback.priority).all()
        for fb in fallbacks:
            cfg = get_llm_config_for_model(fb.fallback_model_id)
            if cfg:
                configs.append(cfg)
    return configs


def _promote_model_to_default(model_id):
    with get_db() as db:
        # Проверяем, существует ли модель и включена ли
        model = db.query(LLMModel).filter(LLMModel.id == model_id, LLMModel.enabled == True).first()
        if not model:
            log_message("LLM_FALLBACK", f"Ошибка: модель {model_id} не найдена или отключена. Нельзя сделать её основной.")
            return

        # Сбрасываем is_default у всех
        db.query(LLMModel).filter(LLMModel.is_default == True).update({LLMModel.is_default: False})
        # Устанавливаем is_default = True для выбранной
        db.query(LLMModel).filter(LLMModel.id == model_id).update({LLMModel.is_default: True})
        db.commit()
        log_message("LLM_FALLBACK", f"Модель {model_id} ({model.model_name}) повышена до основной (is_default=True)")


def get_openai_client():
    global _openai_client
    if _openai_client is None:
        api_key = LLM_TOKEN
        if not api_key:
            api_key = "placeholder"
        _openai_client = OpenAI(
            api_key=api_key,
            base_url="https://bothub.chat/api/v2/openai/v1"
        )
    return _openai_client


def get_default_api_key():
    cfg = get_default_llm_config()
    return cfg.get("api_key") or LLM_TOKEN


def _try_single_llm_call(cfg, messages, timeout):
    import httpx
    ssl_context = ssl.create_default_context()
    ssl_context.check_hostname = False
    ssl_context.verify_mode = ssl.CERT_NONE
    http_client = httpx.Client(verify=ssl_context)
    
    client = OpenAI(
        api_key=cfg.get("api_key") or "placeholder",
        base_url=cfg.get("base_url") or "https://bothub.chat/api/v2/openai/v1",
        http_client=http_client
    )
    return client.chat.completions.create(
        model=cfg.get("model_name") or LLM_MODEL,
        messages=messages,
        temperature=cfg.get("temperature") or 0,
        timeout=cfg.get("timeout") or timeout
    )


def llm_complete_with_config(messages, timeout=120):
    cfg = get_default_llm_config()
    configs = [cfg]
    model_id = cfg.get("model_id")
    if model_id:
        configs += get_fallback_configs(model_id)

    last_error = None
    for attempt, cfg_item in enumerate(configs):
        try:
            response = _try_single_llm_call(cfg_item, messages, timeout)
            
            if attempt > 0 and cfg_item.get("model_id"):
                _promote_model_to_default(cfg_item["model_id"])
            return response
        except Exception as e:
            last_error = e
            if attempt < len(configs) - 1:
                log_message("LLM_FALLBACK", f"{cfg_item.get('model_name')} ошибка: {e}. Пробую fallback... (URL: {cfg_item.get('base_url')}, API_KEY: {bool(cfg_item.get('api_key'))})")
    raise last_error or Exception("Все LLM-модели недоступны")


DOWNLOADS_DIR = os.environ.get("DOWNLOADS_DIR", "app/downloads").strip()
DOWNLOADS_DIR = os.path.normpath(DOWNLOADS_DIR)

try:
    os.makedirs(DOWNLOADS_DIR, exist_ok=True)
    print(f"[DOWNLOADS_DIR] Создана/существует: {DOWNLOADS_DIR}", flush=True)
    test_file = os.path.join(DOWNLOADS_DIR, ".write_test")
    with open(test_file, "w") as f:
        f.write("ok")
    os.remove(test_file)
    print(f"[DOWNLOADS_DIR] Директория доступна для записи", flush=True)
except Exception as e:
    print(f"[DOWNLOADS_DIR] КРИТИЧЕСКАЯ ОШИБКА: не могу создать/писать в {DOWNLOADS_DIR}: {e}", flush=True)
    DOWNLOADS_DIR = "/tmp/downloads"
    os.makedirs(DOWNLOADS_DIR, exist_ok=True)
    print(f"[DOWNLOADS_DIR] Использую альтернативу: {DOWNLOADS_DIR}", flush=True)

# =========================
# DB CONNECTION
# =========================
def _get_db_config():
    with get_db() as db:
        conn = db.query(ConnectionSetting).filter(ConnectionSetting.is_default == True).first()
        if not conn:
            conn = db.query(ConnectionSetting).first()
        if conn:
            return {
                "host": conn.host,
                "database": conn.database_name,
                "user": conn.username,
                "password": conn.password,
                "port": conn.port,
            }
    return {
        "host": "host.docker.internal",
        "database": "dash",
        "user": "postgres",
        "password": "secret",
        "port": 1111,
    }


def get_connection():
    cfg = _get_db_config()
    conn = psycopg2.connect(
        host=cfg["host"],
        database=cfg["database"],
        user=cfg["user"],
        password=cfg["password"],
        port=cfg["port"]
    )
    conn.autocommit = True
    return conn


# =========================
# DB DESCRIPTION (из system_config -> db_desc)
# =========================
def get_db_description():
    desc = get_system_config().get("db_desc", "")
    if desc:
        return desc
    return ""


db_description = get_db_description()


# =========================
# ПРОМПТЫ ИЗ БД
# =========================
def load_prompt_template_from_db(prompt_key):
    with get_db() as db:
        prompt = db.query(Prompt).filter(Prompt.prompt_key == prompt_key).first()
        if not prompt:
            raise FileNotFoundError(f"Prompt not found in DB: {prompt_key}")
        return prompt.content


# =========================
# ПРИМЕНЕНИЕ ШАБЛОНА ПРОМПТА
# =========================
def build_messages_from_template(template_source, placeholders):
    log_message("BUILD_MSGS", f"template_source: {template_source}\nplaceholders keys: {list(placeholders.keys())}")

    template_text = template_source if isinstance(template_source, str) else str(template_source)

    sections = {}
    current_section = None
    current_lines = []

    for line in template_text.split('\n'):
        stripped = line.strip()
        if stripped.startswith('[') and stripped.endswith(']'):
            if current_section:
                sections[current_section] = '\n'.join(current_lines).strip()
            section_name = stripped[1:-1].strip()
            current_section = section_name
            current_lines = []
        else:
            current_lines.append(line)

    if current_section:
        sections[current_section] = '\n'.join(current_lines).strip()

    system_role = sections.get("system_role", "").strip()
    user_role = sections.get("user_role", "").strip()
    task = sections.get("task", "").strip()

    for key, value in placeholders.items():
        val_str = str(value)
        system_role = system_role.replace("{" + key + "}", val_str)
        user_role = user_role.replace("{" + key + "}", val_str)
        task = task.replace("{" + key + "}", val_str)

    if user_role:
        user_content = user_role + "\n\n" + task
    else:
        user_content = task

    user_content = user_content.strip()

    messages = []
    if system_role:
        messages.append({"role": "system", "content": system_role})

    if user_content:
        messages.append({"role": "user", "content": user_content})
    else:
        error_msg = (
            f"Ошибка: содержимое user_content пустое после подстановки. "
            f"Шаблон: {template_source}. "
            f"template keys loaded: {list(sections.keys())}\n"
            f"system_role empty: {len(system_role)==0}\n"
            f"user_role empty: {len(user_role)==0}\n"
            f"task empty: {len(task)==0}"
        )
        log_message("BUILD_MSGS_ERROR", error_msg)
        raise ValueError(error_msg)

    return messages


def build_messages_from_prompt_key(prompt_key, placeholders):
    template_text = load_prompt_template_from_db(prompt_key)
    return build_messages_from_template(template_text, placeholders)


# =========================
# ПОЛЬЗОВАТЕЛИ
# =========================
# В deme у нас пока нет таблицы users, поэтому user_id задаём из session/заголовка
# Для простоты: user_id = int(request.headers.get("X-User-Id", 1))


def get_current_user_id():
    try:
        from flask import request
        return int(request.headers.get("X-User-Id", 1))
    except Exception:
        return 1


# =========================
# GET SCHEMA
# =========================
def get_schema():
    try:
        conn = get_connection()
        cur = conn.cursor()
        query = """
        SELECT table_schema, table_name, column_name, data_type
        FROM information_schema.columns
        WHERE table_schema NOT IN ('pg_catalog', 'information_schema')
        ORDER BY table_schema, table_name, ordinal_position
        """
        cur.execute(query)
        rows = cur.fetchall()
        if not rows:
            raise Exception("empty schema")
        schema = {}
        for schema_name, table, column, dtype in rows:
            full_table = f"{schema_name}.{table}" if schema_name != 'public' else table
            schema.setdefault(full_table, []).append((column, dtype))
        schema_text = ""
        for table, cols in schema.items():
            cols_str = ", ".join([f"{c} ({t})" for c, t in cols])
            schema_text += f"{table}: {cols_str}\n"
        cur.close()
        conn.close()
        return schema_text
    except Exception as e:
        print("SCHEMA ERROR:", str(e), flush=True)
        raise Exception("Не найдена база данных")


# =========================
# PROMPT — ГЕНЕРАЦИЯ SQL
# =========================
def build_prompt(question, schema_text, db_desc):
    examples = load_examples()
    prompt_key = get_system_config().get("prompt_generate_sql_key", "prompt_generate_sql")
    messages = build_messages_from_prompt_key(prompt_key, {
        "SCHEMA_TEXT": schema_text,
        "DB_DESC": db_desc,
        "EXAMPLES": examples,
        "QUESTION": question
    })
    system_role = ""
    user_content = ""
    for msg in messages:
        if msg["role"] == "system":
            system_role = msg["content"]
        elif msg["role"] == "user":
            user_content = msg["content"]
    return {
        "system_role": system_role,
        "user_content": user_content
    }


# =========================
# ЗАГРУЗКА ПРИМЕРОВ (из БД таблицы app.prompts)
# =========================
def load_examples(limit=None):
    """Load examples from DB prompts table with key 'prompt_generate_sql_examples'."""
    if limit is None:
        limit = int(os.environ.get("EXAMPLES", "10"))
    with get_db() as db:
        prompt = db.query(Prompt).filter(
            Prompt.prompt_key == "prompt_generate_sql_examples"
        ).first()
        if not prompt or not prompt.content:
            log_message("BUILD_PROMPT", "EXAMPLES: нет активного промпта prompt_generate_sql_examples")
            return ""
        examples_lines = prompt.content.strip().split("\n")
    examples_text = "\n".join(examples_lines[:limit * 2])
    log_message("BUILD_PROMPT", f"EXAMPLES: загружено {len(examples_lines)} строк (limit={limit})")
    return examples_text


# =========================
# SQL VALIDATOR
# ==========================
def validate_and_fix_sql(question, sql_query, schema_text, db_desc):
    prompt_key = get_system_config().get("prompt_validate_sql_key", "prompt_validate_sql")
    messages = build_messages_from_prompt_key(prompt_key, {
        "SCHEMA_TEXT": schema_text,
        "DB_DESC": db_desc,
        "QUESTION": question,
        "SQL_QUERY": sql_query
    })
    log_llm_request(messages)
    response = llm_complete_with_config(messages, timeout=120)
    validated_sql = response.choices[0].message.content.strip()
    validated_sql = clean_sql(validated_sql)
    return validated_sql


# =========================
# ANALYZE SQL ERROR
# =========================
def analyze_sql_error(sql_query, error_text, schema_text, db_desc):
    prompt_key = get_system_config().get("prompt_analyze_sql_key", "prompt_analyze_sql")
    messages = build_messages_from_prompt_key(prompt_key, {
        "SCHEMA_TEXT": schema_text,
        "DB_DESC": db_desc,
        "SQL_QUERY": sql_query,
        "ERROR_TEXT": error_text
    })
    log_llm_request(messages)
    response = llm_complete_with_config(messages)
    response_text = response.choices[0].message.content.strip()
    return render_markdown(response_text)


# =========================
# EXPLAIN SQL
# =========================
def explain_sql(sql_query, schema_text, db_desc):
    prompt_key = get_system_config().get("prompt_explain_sql_key", "prompt_explain_sql")
    messages = build_messages_from_prompt_key(prompt_key, {
        "SCHEMA_TEXT": schema_text,
        "DB_DESC": db_desc,
        "SQL_QUERY": sql_query
    })
    log_llm_request(messages)
    response = llm_complete_with_config(messages)
    response_text = response.choices[0].message.content.strip()
    return response_text, render_markdown(response_text)


# =========================
# MODIFY SQL FOR BUSINESS TERMS
# =========================
def modify_sql_for_business_terms(sql_query, schema_text, db_desc):
    prompt_key = get_system_config().get("prompt_modify_sql_key", "prompt_modify_sql")
    messages = build_messages_from_prompt_key(prompt_key, {
        "SCHEMA_TEXT": schema_text,
        "DB_DESC": db_desc,
        "SQL_QUERY": sql_query
    })

    messages.append({
        "role": "user",
        "content": (
            "СТРОГОЕ ПРЕДУПРЕЖДЕНИЕ: Ни в коем случае не изменяй формат дат и времени!\n"
            "1. Запрещено использовать TO_CHAR, FORMAT, DATE_FORMAT, CONVERT, AT TIME ZONE\n"
            "   или любые другие функции преобразования для колонок с датами.\n"
            "2. Дата-колонки должны оставаться в исходном виде без каких-либо обёрток.\n"
            "3. Разрешено менять ТОЛЬКО алиасы (AS ...) колонок на бизнес-термины.\n"
            "4. Не добавляй и не удаляй колонки. Не меняй типы данных.\n"
            "5. Сохраняй оригинальный SELECT без изменений структуры."
        )
    })

    log_llm_request(messages)
    response = llm_complete_with_config(messages)
    modified_sql = response.choices[0].message.content.strip()
    modified_sql = clean_sql(modified_sql)
    return modified_sql


# =========================
# ВСПОМОГАТЕЛЬНАЯ: добавить текст с **жирным** и *курсивом*
# =========================
def _add_formatted_text(paragraph, text):
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*)'
    last_end = 0
    for match in re.finditer(pattern, text):
        start, end = match.start(), match.end()
        if start > last_end:
            paragraph.add_run(text[last_end:start])
        if match.group(1).startswith('**'):
            run = paragraph.add_run(match.group(2))
            run.bold = True
        else:
            run = paragraph.add_run(match.group(3))
            run.italic = True
        last_end = end
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


# =========================
# GENERATE WORD REPORT
# =========================
def generate_report(schema_text, db_desc, sql_query, df, filepath, chart_paths=None, prompt_id=None):
    current_date = datetime.datetime.now().strftime("%d.%m.%Y")
    current_date_full = datetime.datetime.now().strftime("%d.%m.%Y %H:%M")

    placeholders = {
        "SCHEMA_TEXT": schema_text,
        "DB_DESC": db_desc,
        "SQL_QUERY": sql_query,
        "DATA_STR": "",
        "CURRENT_DATE": current_date,
        "CURRENT_DATE_FULL": current_date_full,
    }

    # Формируем DATA_STR как табличный текст
    if df is not None and not df.empty:
        data_lines = [" | ".join(str(c) for c in df.columns)]
        data_lines.append("-+-".join(["-" * min(len(str(c)), 20) for c in df.columns]))
        for _, row in df.head(100).iterrows():
            data_lines.append(" | ".join(str(v) for v in row))
        placeholders["DATA_STR"] = "\n".join(data_lines)

    if prompt_id:
        from models import PromptReport
        with get_db() as db:
            pr = db.query(PromptReport).filter(PromptReport.id == prompt_id).first()
        if pr:
            system_text = pr.content
            for k, v in placeholders.items():
                system_text = system_text.replace("{" + k + "}", str(v))
            messages = [{"role": "system", "content": system_text}]
        else:
            prompt_key = get_system_config().get("prompt_report_key", "prompt_report")
            messages = build_messages_from_prompt_key(prompt_key, placeholders)
    else:
        prompt_key = get_system_config().get("prompt_report_key", "prompt_report")
        messages = build_messages_from_prompt_key(prompt_key, placeholders)

    messages.append({
        "role": "user",
        "content": (
            f"СТРОГОЕ ПРЕДУПРЕЖДЕНИЕ О ДАТАХ:\n"
            f"Сегодня {current_date} ({current_date_full}).\n"
            f"Используй ТОЛЬКО эту дату как 'сегодня', 'текущая дата', 'дата отчёта'.\n"
            f"1. Никогда не выдумывай и не изменяй даты.\n"
            f"2. Если нужно указать дату отчёта — пиши '{current_date}'.\n"
            f"3. Даты из данных таблицы бери ТОЛЬКО из колонок с данными.\n"
            f"4. Не заменяй даты из данных на другие.\n"
            f"5. Все числовые значения должны быть взяты непосредственно из данных.\n"
            f"6. Если в данных нет какой-то информации — не добавляй её от себя.\n"
            f"7. Пиши отчёт строго по фактам из предоставленной таблицы."
        )
    })

    log_llm_request(messages)
    response = llm_complete_with_config(messages)
    report_text = response.choices[0].message.content.strip()

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[REPORT] python-docx не установлен. Установите: pip install python-docx", flush=True)
        raise Exception("python-docx не установлен. Выполните: pip install python-docx")

    doc = Document()
    doc.add_heading("Аналитический отчет", level=1)

    p = doc.add_paragraph()
    run = p.add_run(f"Дата формирования: {current_date_full}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    p.space_after = Pt(12)

    for line in report_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[2:])
        elif line.startswith('* ') or line.startswith('+ '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

    if chart_paths:
        doc.add_page_break()
        doc.add_heading("Графики", level=2)
        for i, chart_path in enumerate(chart_paths):
            if os.path.exists(chart_path):
                try:
                    doc.add_paragraph(f"График {i+1}").runs[0].bold = True
                    doc.add_picture(chart_path, width=Inches(6.0))
                    doc.add_paragraph("")
                except Exception as e:
                    print(f"[REPORT] Ошибка вставки графика {i+1}: {e}", flush=True)

    if df is not None and not df.empty:
        doc.add_page_break()
        doc.add_heading("Результат выполнения SQL", level=2)
        num_cols = len(df.columns)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        display_rows = df.head(100)
        for _, row_data in display_rows.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = str(val) if val is not None else "NULL"

        if len(df) > 100:
            p = doc.add_paragraph()
            run = p.add_run(f"Показано 100 из {len(df)} строк")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"Всего строк: {len(df)}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(filepath)
    print(f"[REPORT] Сохранён: {filepath}", flush=True)


# =========================
# SQL UTILS
# =========================
def clean_sql(sql):
    sql = re.sub(r"```sql", "", sql, flags=re.IGNORECASE)
    sql = re.sub(r"```", "", sql)
    sql = sql.strip()
    match = re.search(r"(select[\s\S]*)", sql, re.IGNORECASE)
    if match:
        sql = match.group(1)
    return sql.strip().rstrip(";")


def fix_limit(sql):
    if "limit" not in sql.lower():
        sql += f" LIMIT {QUERY_LIMIT}"
    return sql


def validate_sql(sql):
    banned = ["drop", "delete", "update", "insert", "alter", "truncate", "grant", "revoke"]
    if not sql.lower().startswith("select"):
        raise Exception("Разрешены только SELECT")
    for b in banned:
        if b in sql.lower():
            raise Exception(f"Запрещено: {b}")
    return sql


def run_sql_to_df(sql):
    conn = get_connection()
    df = pd.read_sql_query(sql, conn)
    conn.close()
    return df


def sql_to_one_line(sql):
    sql = clean_sql(sql)
    sql = re.sub(r"\s+", " ", sql)
    return sql.strip().rstrip(";")


# =========================
# SAVE EXCEL
# =========================
def save_excel(df, filename=None):
    if filename is None:
        filename = "Chart.xlsx"

    df_export = df.copy()
    for col in df_export.columns:
        if 'datetime' in str(df_export[col].dtype):
            if df_export[col].dt.tz is not None:
                df_export[col] = df_export[col].dt.tz_localize(None)
            df_export[col] = df_export[col].astype(str)

    excel_path = os.path.join(DOWNLOADS_DIR, filename)
    with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
        df_export.to_excel(writer, index=False, sheet_name="Data")
        ws = writer.sheets["Data"]
        for column_cells in ws.columns:
            length = max(
                len(str(cell.value)) if cell.value is not None else 0
                for cell in column_cells
            )
            adjusted_width = min(length + 5, 50)
            ws.column_dimensions[column_cells[0].column_letter].width = adjusted_width
    print(f"[EXCEL] {excel_path}", flush=True)


# =========================
# CHARTS
# =========================
def build_charts(df, color=None):
    figs = []
    if df.empty:
        return figs
    cols = list(df.columns)
    if len(cols) == 1:
        fig = px.histogram(df, x=cols[0], title=cols[0])
        if color:
            fig.update_traces(marker_color=color)
        figs.append({"title": cols[0], "fig": fig})
        return figs
    if len(cols) == 2:
        fig = px.bar(df, x=cols[0], y=cols[1], title=f"{cols[0]} → {cols[1]}")
        if color:
            fig.update_traces(marker_color=color)
        figs.append({"title": f"{cols[0]}_{cols[1]}", "fig": fig})
        return figs
    x_col = cols[0]
    for y_col in cols[1:]:
        fig = px.bar(df, x=x_col, y=y_col, title=f"{x_col} → {y_col}")
        if color:
            fig.update_traces(marker_color=color)
        figs.append({"title": f"{x_col}_{y_col}", "fig": fig})
    return figs


def build_chart_body_html(figs):
    html_parts = []
    for i, item in enumerate(figs):
        fig = item["fig"]
        title = item["title"]
        html_parts.append(f"""
            <h2>{title}</h2>
            {fig.to_html(full_html=False, include_plotlyjs='cdn')}
            <hr>
        """)
    return "".join(html_parts)


def create_jpg_collage(image_paths, output_path):
    images = [Image.open(p) for p in image_paths]
    widths = [img.width for img in images]
    heights = [img.height for img in images]
    max_width = max(widths)
    total_height = sum(heights)
    collage = Image.new("RGB", (max_width, total_height), color=(30, 30, 30))
    y_offset = 0
    for img in images:
        collage.paste(img, (0, y_offset))
        y_offset += img.height
    collage.save(output_path, "JPEG")


def build_chart_body_only(figs):
    html_parts = []
    for i, item in enumerate(figs):
        fig = item["fig"]
        title = item["title"]
        html_parts.append(f"""
            <h2>{title}</h2>
            {fig.to_html(full_html=False, include_plotlyjs='cdn')}
            <hr>
        """)
    chart_body = "".join(html_parts)
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Charts</title>
        <style>
            body {{
                background:#1e1e1e;
                color:white;
                font-family:Arial;
                padding:20px;
            }}
        </style>
    </head>
    <body>
        {chart_body}
    </body>
    </html>
    """
    return {"full_html": full_html, "chart_body": chart_body}


def save_chart_outputs(figs, base_dir):
    os.makedirs(base_dir, exist_ok=True)
    html_parts = []
    jpg_paths = []
    for i, item in enumerate(figs):
        fig = item["fig"]
        title = item["title"]
        html_parts.append(f"""
            <h2>{title}</h2>
            {fig.to_html(full_html=False, include_plotlyjs='cdn')}
            <hr>
        """)
        jpg_path = os.path.join(base_dir, f"chart_{i+1}.jpg")
        fig.write_image(jpg_path, format="jpg", width=1400, height=900, scale=2)
        jpg_paths.append(jpg_path)
    chart_body = "".join(html_parts)
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Charts</title>
        <style>
            body {{
                background:#1e1e1e;
                color:white;
                font-family:Arial;
                padding:20px;
            }}
        </style>
    </head>
    <body>
        {chart_body}
    </body>
    </html>
    """
    html_path = os.path.join(base_dir, "chart.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(full_html)
    collage_path = os.path.join(base_dir, "chart.jpg")
    create_jpg_collage(jpg_paths, collage_path)
    return {"full_html": full_html, "chart_body": chart_body}


# =========================
# ЛОГИРОВАНИЕ (универсальное)
# =========================
def _write_log(log_entry: str):
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    log_path = os.path.join(DOWNLOADS_DIR, f"llm_{date_str}.txt")
    try:
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(log_entry)
    except Exception as e:
        print(f"[LOG ERROR] Не удалось записать лог в {log_path}: {e}", flush=True)


def log_message(tag: str, text: str):
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    entry = f"[{now}] [{tag}]\n{text}\n{'='*60}\n"
    _write_log(entry)
    print(f"[{tag}] {text[:200]}...", flush=True)


def log_llm_request(messages):
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    prompt_text = ""
    for msg in messages:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        prompt_text += f"[{role}]\n{content}\n\n"
    entry = (
        f"Дата и время запроса ({now}):\n"
        f"{prompt_text}"
        f"=============================================\n"
    )
    _write_log(entry)
    print(f"[LLM LOG] Запрос записан в {DOWNLOADS_DIR}", flush=True)


def log_user_llm_request(user, question, llm_answer):
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    log_path = os.path.join(DOWNLOADS_DIR, f"req_{date_str}.txt")
    entry = (
        f"Пользователь: {user}\n"
        f"Дата и время: {now}\n"
        f"Вопрос: {question}\n"
        f"Ответ LLM: {llm_answer}\n"
        "----------------\n"
    )
    print(f"[LOG_USER_LLM] Attempting to write to: {log_path}")
    try:
        os.makedirs(DOWNLOADS_DIR, exist_ok=True)
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(entry)
        print(f"[LOG_USER_LLM] Written to: {log_path}")
    except Exception as e:
        print(f"[LOG_USER_LLM] ERROR: {e}", flush=True)

    # Дополнительно: пишем в llm_*.txt для видимости
    llm_log_path = os.path.join(DOWNLOADS_DIR, f"llm_{date_str}.txt")
    _write_log(entry)
    _write_log(entry)
    llm_entry = f"[USER REQUEST] {user} | {question}\n[LLM RESPONSE] {llm_answer}\n{'='*60}\n"
    try:
        with open(llm_log_path, "a", encoding="utf-8") as f:
            f.write(llm_entry)
    except Exception as e:
        print(f"[LOG_USER_LLM] ERROR writing to llm log: {e}", flush=True)


def log_request_response(tag, request_text, response_text):
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    log_path = os.path.join(DOWNLOADS_DIR, f"requests_{date_str}.txt")
    try:
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"Время: {now}\n")
            f.write(f"Запрос: {request_text}\n")
            f.write(f"Ответ: {response_text}\n\n")
            f.write(f"===============================================================\n\n")
    except Exception as e:
        print(f"[REQUESTS LOG ERROR] {e}", flush=True)
    print(f"[REQUESTS LOG] [{tag}] Запись добавлена в {log_path}", flush=True)


# =========================
# MARKDOWN RENDERING
# =========================
def render_markdown(text):
    if not text:
        return ""
    return markdown.markdown(
        text,
        extensions=["fenced_code", "tables", "codehilite", "nl2br"]
    )


# =========================
# DATED FILENAME
# =========================
def get_dated_filename(prefix, ext, label=""):
    safe_label = re.sub(r'[^a-zA-Zа-яА-Я0-9_]', '_', label)[:30] if label else ""
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    name = f"{prefix}_{safe_label}_{ts}.{ext}" if safe_label else f"{prefix}_{ts}.{ext}"
    return name
